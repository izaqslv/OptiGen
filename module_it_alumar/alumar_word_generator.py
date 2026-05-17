import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

class AlumarWordGenerator:
    """
    Gerador Profissional de Word no padrão Alumar com suporte a Imagens Ilustrativas.
    """
    def __init__(self, data: dict, output_path: str):
        self.data = data
        self.output_path = output_path
        self.doc = Document()
        
    def set_cell_background(self, cell, fill):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def generate(self):
        section = self.doc.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)

        # 2. CABEÇALHO
        header_table = self.doc.add_table(rows=3, cols=3)
        header_table.style = 'Table Grid'
        header_table.autofit = False
        
        widths = [Inches(1.2), Inches(3.8), Inches(5.5)]
        for row in header_table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = width

        cell_alumar = header_table.cell(0, 0)
        cell_alumar.text = "ALUMAR"
        self.set_cell_background(cell_alumar, "D3D3D3")
        run = cell_alumar.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(12)
        cell_alumar.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        title_cell = header_table.cell(0, 1)
        title_cell.merge(header_table.cell(0, 2))
        title_cell.text = f"Instrução de Trabalho - {self.data.get('titulo', '')}"
        run = title_cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(11)

        header_table.cell(1, 0).text = f"Nº DA IT: {self.data.get('id', 'NOVO')}"
        header_table.cell(1, 1).text = f"LOCAL: {self.data.get('local', '')}"
        
        # Objetivo com quebra de linha automática
        obj_cell = header_table.cell(1, 2)
        obj_cell.text = f"OBJETIVO: {self.data.get('objetivo', '')}"

        header_table.cell(2, 0).text = f"VERSÃO: {self.data.get('versao', '001')}"
        header_table.cell(2, 1).text = f"AUTOR: {self.data.get('autor', 'NewGen/OptiGen AI')}"
        header_table.cell(2, 2).text = "APROVADOR: Pendente"

        for row in header_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        self.doc.add_paragraph()

        # 3. MATRIZ DE SEGURANÇA
        safety_table = self.doc.add_table(rows=2, cols=4)
        safety_table.style = 'Table Grid'
        headers = ["RISCOS", "CONTROLES CRÍTICOS", "CRITÉRIOS DE PARADA", "EQUIPAMENTOS / EPIS"]
        bg_colors = ["FF0000", "008000", "FFA500", "D3D3D3"]
        
        for i, (header, color) in enumerate(zip(headers, bg_colors)):
            cell = safety_table.cell(0, i)
            cell.text = header
            self.set_cell_background(cell, color)
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255) if i < 3 else RGBColor(0, 0, 0)
            run.font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        ms = self.data.get('matriz_seguranca', {})
        safety_table.cell(1, 0).text = "\n".join([f"• {r}" for r in ms.get('riscos', [])])
        safety_table.cell(1, 1).text = "\n".join([f"• {c}" for c in ms.get('controles_criticos', [])])
        safety_table.cell(1, 2).text = "\n".join([f"• {p}" for p in ms.get('criterios_parada', [])])
        lista_equip = ms.get('equipamentos_ferramentas', []) + ms.get('epis', [])
        safety_table.cell(1, 3).text = "\n".join([f"• {e}" for e in lista_equip])

        for cell in safety_table.rows[1].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

        self.doc.add_paragraph()

        # 4. FLUXO DE EXECUÇÃO COM COLUNA DE FOTO
        exec_table = self.doc.add_table(rows=1, cols=5)
        exec_table.style = 'Table Grid'
        hdr_cells = exec_table.rows[0].cells
        headers_exec = ['PASSO', 'O QUE FAZER?', 'COMO FAZER?', 'SEGURANÇA', 'FOTO ILUSTRATIVA']
        
        for i, h in enumerate(headers_exec):
            hdr_cells[i].text = h
            self.set_cell_background(hdr_cells[i], "000080")
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for step in self.data.get('fluxo_execucao', []):
            row = exec_table.add_row()
            row.height = Inches(1.5) # Garante espaço para a foto
            row_cells = row.cells
            row_cells[0].text = str(step.get('passo_n', ''))
            row_cells[1].text = step.get('o_que_fazer', '')
            row_cells[2].text = f"{step.get('como_fazer', '')}\n\nMotivo: {step.get('por_que_fazer', '')}"
            row_cells[3].text = "\n".join(step.get('medidas_controle', []))
            row_cells[4].text = "[COLE A FOTO AQUI]" # Placeholder visual
            
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        self.doc.save(self.output_path)
        return self.output_path

if __name__ == "__main__":
    dados_exemplo = {
        "titulo": "Operação de Ponte Rolante ECL",
        "local": "Redução sala de cubas L-3",
        "objetivo": "Realizar a troca periódica dos ânodos e manutenção preventiva.",
        "matriz_seguranca": {"riscos": ["Risco 1"], "controles_criticos": ["Controle 1"], "criterios_parada": ["Parada 1"], "equipamentos_ferramentas": ["Equip 1"]},
        "fluxo_execucao": [{"passo_n": 1, "o_que_fazer": "Fazer algo", "como_fazer": "Assim", "por_que_fazer": "Pois sim", "medidas_controle": ["EPI"]}]
    }
    gen = AlumarWordGenerator(dados_exemplo, "Teste_IT_Word_Melhorada.docx")
    gen.generate()
    print("Word Melhorado gerado com sucesso!")
